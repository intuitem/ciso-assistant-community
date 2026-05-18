"""
Document text extraction and chunking for RAG ingestion.
Each extractor handles a specific file type and produces semantically meaningful chunks.
"""

import csv
import io
import re
import structlog
from dataclasses import dataclass

logger = structlog.get_logger(__name__)

CHUNK_SIZE = 500  # Target tokens per chunk (approx 4 chars per token)
CHUNK_OVERLAP = 50  # Token overlap between chunks
MAX_CHUNK_CHARS = CHUNK_SIZE * 4

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")

# Lazy state for Docling. The converter loads ~300 MB of models on first
# instantiation, so we keep it process-global. `_docling_unavailable` flips
# permanently after import or first-init failure so we don't retry every call.
_docling_converter = None
_docling_unavailable = False


def _split_long_text(text: str, max_len: int) -> list[str]:
    """Split text that exceeds max_len on sentence boundaries."""
    parts = []
    while len(text) > max_len:
        # Try to split on sentence boundary
        cut = text.rfind(". ", 0, max_len)
        if cut < max_len // 2:
            cut = text.rfind(" ", 0, max_len)
        if cut < max_len // 2:
            cut = max_len
        parts.append(text[: cut + 1].strip())
        text = text[cut + 1 :].strip()
    if text:
        parts.append(text)
    return parts


@dataclass
class Chunk:
    """A chunk of text extracted from a document."""

    text: str
    index: int
    metadata: dict


def get_extractor(content_type: str):
    """Return the appropriate extractor for the given content type."""
    extractors = {
        "application/pdf": extract_pdf,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": extract_excel,
        # .xls (vnd.ms-excel) excluded — openpyxl only supports .xlsx
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx,
        # Markdown is plain text once the LLM sees it — no need for a real
        # parser. Same chunker as .txt; headings stay inline as `#` lines.
        "text/markdown": extract_text,
        "text/csv": extract_csv,
        "text/plain": extract_text,
    }
    return extractors.get(content_type)


def _get_docling_converter():
    """Lazy-init the Docling converter. Returns None if unavailable.

    First call pays the model-load cost (~300 MB downloaded once, cached in
    ~/.cache/docling on subsequent runs). Failures flip a permanent flag so
    we skip retrying on every extraction.
    """
    global _docling_converter, _docling_unavailable
    if _docling_unavailable:
        return None
    if _docling_converter is not None:
        return _docling_converter
    try:
        from docling.document_converter import DocumentConverter

        _docling_converter = DocumentConverter()
        logger.info("docling_initialized")
        return _docling_converter
    except Exception as e:
        logger.warning("docling_unavailable", error=str(e))
        _docling_unavailable = True
        return None


def _extract_with_docling(file, hint: str) -> list[Chunk] | None:
    """Try Docling for layout-aware extraction. Returns None on failure so the
    caller can fall back to its prior implementation.
    """
    converter = _get_docling_converter()
    if converter is None:
        return None
    try:
        from docling.datamodel.base_models import DocumentStream

        file.seek(0)
        data = file.read()
        source = DocumentStream(name=hint, stream=io.BytesIO(data))
        result = converter.convert(source)
        markdown = result.document.export_to_markdown()
        if not markdown.strip():
            return None
        return _markdown_to_chunks(markdown, base_metadata={"extractor": "docling"})
    except Exception as e:
        logger.warning("docling_extract_failed", error=str(e), hint=hint)
        return None


def _markdown_to_chunks(markdown: str, base_metadata: dict) -> list[Chunk]:
    """Chunk markdown while prefixing each chunk with its heading breadcrumb.

    The LLM sees something like:
        [5. Access Control > 5.2 Authentication]
        MFA shall be enforced for all administrative accounts. ...

    so a control extracted from this chunk can be tagged with its originating
    section. Tables are emitted by Docling as markdown pipe tables and stay
    inline — they don't need special handling.
    """
    chunks: list[Chunk] = []
    heading_stack: list[str] = []
    current_lines: list[str] = []
    current_len = 0

    def flush():
        nonlocal current_lines, current_len
        if not current_lines:
            return
        body = "\n".join(current_lines).strip()
        current_lines = []
        current_len = 0
        if not body:
            return
        breadcrumb = " > ".join(filter(None, heading_stack))
        prefix = f"[{breadcrumb}]\n" if breadcrumb else ""
        text = (prefix + body).strip()
        meta = dict(base_metadata)
        if breadcrumb:
            meta["section"] = breadcrumb
        parts = (
            _split_long_text(text, MAX_CHUNK_CHARS)
            if len(text) > MAX_CHUNK_CHARS
            else [text]
        )
        for part in parts:
            chunks.append(Chunk(text=part, index=len(chunks), metadata=dict(meta)))

    for line in markdown.split("\n"):
        m = _HEADING_RE.match(line)
        if m:
            flush()
            level = len(m.group(1))
            title = m.group(2).strip()
            # Trim the stack to one slot above the new heading's level, then push
            heading_stack[:] = heading_stack[: level - 1]
            while len(heading_stack) < level - 1:
                heading_stack.append("")
            heading_stack.append(title)
            continue
        current_lines.append(line)
        current_len += len(line) + 1
        if current_len >= MAX_CHUNK_CHARS:
            flush()

    flush()
    return chunks


def extract_pdf(file) -> list[Chunk]:
    """Extract text from PDF.

    Prefers Docling for layout-aware parsing (preserves tables and section
    headings). Falls back to PyMuPDF if Docling is unavailable, fails, or
    returns nothing useful — preserves the prior behavior on edge-case files.
    """
    chunks = _extract_with_docling(file, hint="document.pdf")
    if chunks:
        return chunks

    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF not installed, cannot extract PDF")
        return []

    chunks = []
    file.seek(0)
    doc = fitz.open(stream=file.read(), filetype="pdf")

    for page_num, page in enumerate(doc):
        text = page.get_text().strip()
        if not text:
            continue

        # Split into paragraph-sized chunks
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        current_chunk = ""

        for para in paragraphs:
            # Split oversized paragraphs to stay within chunk limits
            if len(para) > CHUNK_SIZE * 4:
                para_parts = _split_long_text(para, CHUNK_SIZE * 4)
            else:
                para_parts = [para]

            for part in para_parts:
                if len(current_chunk) + len(part) > CHUNK_SIZE * 4:
                    if current_chunk:
                        chunks.append(
                            Chunk(
                                text=current_chunk.strip(),
                                index=len(chunks),
                                metadata={"page": page_num + 1},
                            )
                        )
                    current_chunk = part
                else:
                    current_chunk += "\n\n" + part if current_chunk else part

        if current_chunk:
            chunks.append(
                Chunk(
                    text=current_chunk.strip(),
                    index=len(chunks),
                    metadata={"page": page_num + 1},
                )
            )

    doc.close()
    return chunks


def extract_excel(file) -> list[Chunk]:
    """
    Extract from Excel preserving column semantics.
    Each row becomes a chunk with column headers as context.
    """
    try:
        import openpyxl
    except ImportError:
        logger.warning("openpyxl not installed, cannot extract Excel")
        return []

    chunks = []
    file.seek(0)
    wb = openpyxl.load_workbook(file, read_only=True, data_only=True)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # First row as headers
        headers = [str(h) if h else f"Column_{i}" for i, h in enumerate(rows[0])]

        # Each row becomes a chunk
        for row_idx, row in enumerate(rows[1:], start=2):
            parts = []
            for header, value in zip(headers, row):
                if value is not None and str(value).strip():
                    parts.append(f"{header}: {value}")

            if parts:
                text = f"Sheet: {sheet_name} | Row {row_idx}\n" + " | ".join(parts)
                chunks.append(
                    Chunk(
                        text=text,
                        index=len(chunks),
                        metadata={"sheet": sheet_name, "row": row_idx},
                    )
                )

    wb.close()
    return chunks


def extract_csv(file) -> list[Chunk]:
    """Extract from CSV with column headers as context."""
    chunks = []
    file.seek(0)

    content = file.read()
    if isinstance(content, bytes):
        content = content.decode("utf-8", errors="replace")

    reader = csv.reader(io.StringIO(content))
    headers = next(reader, None)
    if not headers:
        return []

    for row_idx, row in enumerate(reader, start=2):
        parts = []
        for header, value in zip(headers, row):
            if value and value.strip():
                parts.append(f"{header}: {value}")

        if parts:
            text = " | ".join(parts)
            chunks.append(
                Chunk(
                    text=text,
                    index=len(chunks),
                    metadata={"row": row_idx},
                )
            )

    return chunks


def extract_docx(file) -> list[Chunk]:
    """Extract from .docx.

    Prefers Docling — it walks the document tree, preserves the order of
    paragraphs and tables, and emits markdown with heading hierarchy intact.
    Falls back to the python-docx implementation (paragraphs first, tables
    tacked on at the end) if Docling fails.
    """
    chunks = _extract_with_docling(file, hint="document.docx")
    if chunks:
        return chunks

    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed, cannot extract .docx")
        return []

    file.seek(0)
    doc = Document(file)

    text_pieces: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_pieces.append(t)

    for table_idx, table in enumerate(doc.tables, start=1):
        rendered_rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                rendered_rows.append("| " + " | ".join(cells) + " |")
        if rendered_rows:
            text_pieces.append(f"[Table {table_idx}]\n" + "\n".join(rendered_rows))

    if not text_pieces:
        return []

    return _paragraphs_to_chunks(text_pieces, base_metadata={})


def _paragraphs_to_chunks(paragraphs: list[str], base_metadata: dict) -> list[Chunk]:
    """Greedy merge of paragraphs into ~MAX_CHUNK_CHARS-sized chunks.

    Factored out so .docx and .txt share one chunker — the only difference
    between them was where the paragraphs come from.
    """
    chunks: list[Chunk] = []
    current_chunk = ""

    for para in paragraphs:
        parts = (
            _split_long_text(para, MAX_CHUNK_CHARS)
            if len(para) > MAX_CHUNK_CHARS
            else [para]
        )
        for part in parts:
            if len(current_chunk) + len(part) > MAX_CHUNK_CHARS:
                if current_chunk:
                    chunks.append(
                        Chunk(
                            text=current_chunk.strip(),
                            index=len(chunks),
                            metadata=dict(base_metadata),
                        )
                    )
                current_chunk = part
            else:
                current_chunk += "\n\n" + part if current_chunk else part

    if current_chunk:
        chunks.append(
            Chunk(
                text=current_chunk.strip(),
                index=len(chunks),
                metadata=dict(base_metadata),
            )
        )

    return chunks


def extract_text(file) -> list[Chunk]:
    """Extract plain text (or markdown), splitting into chunks by paragraph."""
    file.seek(0)
    content = file.read()
    if isinstance(content, bytes):
        content = content.decode("utf-8", errors="replace")

    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    return _paragraphs_to_chunks(paragraphs, base_metadata={})
