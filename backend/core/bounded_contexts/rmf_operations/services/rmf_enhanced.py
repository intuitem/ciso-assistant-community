"""
RMF Enhanced Services

Enhanced RMF operations inspired by OpenRMF patterns:
- SCAP result import (OpenSCAP, SCC)
- RMF document generation (POA&M, Test Plan, RAR)
- System scoring dashboard
- Enhanced Nessus correlation
- Asset metadata management
"""

import json
import logging
import uuid
import xml.etree.ElementTree as ET
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime, date
from io import BytesIO

logger = logging.getLogger(__name__)


@dataclass
class ScapResult:
    """Parsed SCAP scan result"""
    benchmark_id: str
    benchmark_title: str
    benchmark_version: str
    profile_id: str
    target_id: str
    target_hostname: str
    start_time: datetime
    end_time: datetime
    rules: List[Dict[str, Any]] = field(default_factory=list)
    score: float = 0.0
    pass_count: int = 0
    fail_count: int = 0
    error_count: int = 0
    notchecked_count: int = 0


@dataclass
class AssetMetadata:
    """Enhanced asset metadata (OpenRMF pattern)"""
    hostname: str
    ip_addresses: List[str] = field(default_factory=list)
    mac_addresses: List[str] = field(default_factory=list)
    fqdn: str = ""
    technology_area: str = ""
    asset_type: str = "computing"
    role: str = ""
    operating_system: str = ""
    os_version: str = ""
    location: str = ""
    department: str = ""
    system_administrator: str = ""
    is_internet_facing: bool = False
    classification: str = "unclassified"


@dataclass
class RMFDocumentResult:
    """Result of RMF document generation"""
    success: bool
    document_type: str
    content: bytes
    filename: str
    errors: List[str] = field(default_factory=list)


@dataclass
class SystemScore:
    """Aggregate system compliance score"""
    system_group_id: str
    system_name: str
    total_checklists: int
    total_findings: int
    cat1_open: int
    cat1_closed: int
    cat2_open: int
    cat2_closed: int
    cat3_open: int
    cat3_closed: int
    compliance_percentage: float
    risk_score: float
    last_calculated: datetime


class SCAPParser:
    """
    Parser for SCAP (Security Content Automation Protocol) results.

    Supports:
    - XCCDF 1.2 results (OpenSCAP, SCC)
    - OVAL results
    - ARF (Assessment Result Format)
    """

    XCCDF_NS = {
        'xccdf12': 'http://checklists.nist.gov/xccdf/1.2',
        'xccdf11': 'http://checklists.nist.gov/xccdf/1.1',
        'oval': 'http://oval.mitre.org/XMLSchema/oval-results-5',
        'arf': 'http://scap.nist.gov/schema/asset-reporting-format/1.1',
        'ai': 'http://scap.nist.gov/schema/asset-identification/1.1',
    }

    RESULT_MAP = {
        'pass': 'not_a_finding',
        'fail': 'open',
        'error': 'not_reviewed',
        'unknown': 'not_reviewed',
        'notapplicable': 'not_applicable',
        'notchecked': 'not_reviewed',
        'notselected': 'not_applicable',
        'informational': 'not_a_finding',
        'fixed': 'not_a_finding',
    }

    def parse_xccdf_results(self, xml_content: str) -> ScapResult:
        """
        Parse XCCDF 1.2 test results.

        Args:
            xml_content: XCCDF result XML content

        Returns:
            ScapResult with parsed data
        """
        root = ET.fromstring(xml_content)

        # Detect XCCDF version
        if root.tag.endswith('Benchmark'):
            # This is a benchmark, not results - need TestResult
            return self._parse_xccdf_benchmark_with_results(root)
        elif root.tag.endswith('TestResult'):
            return self._parse_xccdf_test_result(root)
        elif 'asset-report-collection' in root.tag:
            return self._parse_arf_results(root)
        else:
            raise ValueError(f"Unknown XCCDF format: {root.tag}")

    def _parse_xccdf_test_result(self, root: ET.Element) -> ScapResult:
        """Parse XCCDF TestResult element"""
        ns = self.XCCDF_NS

        # Try XCCDF 1.2 first, then 1.1
        for ns_prefix in ['xccdf12', 'xccdf11']:
            benchmark = root.find(f'.//{{{ns[ns_prefix]}}}benchmark', ns)
            if benchmark is not None:
                break

        result = ScapResult(
            benchmark_id=root.get('id', ''),
            benchmark_title='',
            benchmark_version='',
            profile_id='',
            target_id='',
            target_hostname='',
            start_time=datetime.now(),
            end_time=datetime.now(),
        )

        # Extract benchmark info
        if benchmark is not None:
            result.benchmark_id = benchmark.get('href', benchmark.get('idref', ''))

        # Extract target info
        target = root.find('.//{http://checklists.nist.gov/xccdf/1.2}target', ns)
        if target is not None and target.text:
            result.target_hostname = target.text

        # Extract times
        start_time = root.get('start-time')
        end_time = root.get('end-time')
        if start_time:
            try:
                result.start_time = datetime.fromisoformat(start_time.replace('Z', '+00:00'))
            except ValueError:
                pass
        if end_time:
            try:
                result.end_time = datetime.fromisoformat(end_time.replace('Z', '+00:00'))
            except ValueError:
                pass

        # Extract rule results
        rules = []
        for rule_result in root.findall('.//{http://checklists.nist.gov/xccdf/1.2}rule-result', ns):
            rule_id = rule_result.get('idref', '')

            result_elem = rule_result.find('{http://checklists.nist.gov/xccdf/1.2}result', ns)
            rule_status = result_elem.text if result_elem is not None else 'unknown'

            # Map to STIG status
            mapped_status = self.RESULT_MAP.get(rule_status.lower(), 'not_reviewed')

            # Extract severity from rule-result or from original rule
            severity = rule_result.get('severity', 'medium')

            rules.append({
                'rule_id': rule_id,
                'result': rule_status,
                'mapped_status': mapped_status,
                'severity': severity,
            })

            # Update counts
            if rule_status.lower() == 'pass':
                result.pass_count += 1
            elif rule_status.lower() == 'fail':
                result.fail_count += 1
            elif rule_status.lower() == 'error':
                result.error_count += 1
            else:
                result.notchecked_count += 1

        result.rules = rules

        # Calculate score
        total = result.pass_count + result.fail_count
        if total > 0:
            result.score = (result.pass_count / total) * 100

        return result

    def _parse_xccdf_benchmark_with_results(self, root: ET.Element) -> ScapResult:
        """Parse XCCDF Benchmark that contains TestResult"""
        ns = self.XCCDF_NS

        # Get benchmark info
        benchmark_id = root.get('id', '')

        title_elem = root.find('.//{http://checklists.nist.gov/xccdf/1.2}title', ns)
        benchmark_title = title_elem.text if title_elem is not None else ''

        version_elem = root.find('.//{http://checklists.nist.gov/xccdf/1.2}version', ns)
        benchmark_version = version_elem.text if version_elem is not None else ''

        # Find TestResult
        test_result = root.find('.//{http://checklists.nist.gov/xccdf/1.2}TestResult', ns)
        if test_result is None:
            test_result = root.find('.//{http://checklists.nist.gov/xccdf/1.1}TestResult', ns)

        if test_result is not None:
            result = self._parse_xccdf_test_result(test_result)
            result.benchmark_id = benchmark_id
            result.benchmark_title = benchmark_title
            result.benchmark_version = benchmark_version
            return result

        # No TestResult found - return empty result
        return ScapResult(
            benchmark_id=benchmark_id,
            benchmark_title=benchmark_title,
            benchmark_version=benchmark_version,
            profile_id='',
            target_id='',
            target_hostname='',
            start_time=datetime.now(),
            end_time=datetime.now(),
        )

    def _parse_arf_results(self, root: ET.Element) -> ScapResult:
        """Parse ARF (Assessment Result Format) results"""
        ns = self.XCCDF_NS

        # ARF contains assets and reports
        result = ScapResult(
            benchmark_id='',
            benchmark_title='',
            benchmark_version='',
            profile_id='',
            target_id='',
            target_hostname='',
            start_time=datetime.now(),
            end_time=datetime.now(),
        )

        # Find XCCDF results within ARF
        xccdf_result = root.find('.//{http://checklists.nist.gov/xccdf/1.2}TestResult', ns)
        if xccdf_result is not None:
            return self._parse_xccdf_test_result(xccdf_result)

        return result

    def convert_to_stig_findings(
        self,
        scap_result: ScapResult,
        cci_mapping: Dict[str, List[str]]
    ) -> List[Dict[str, Any]]:
        """
        Convert SCAP results to STIG-compatible vulnerability findings.

        Args:
            scap_result: Parsed SCAP result
            cci_mapping: Mapping of rule IDs to CCI IDs

        Returns:
            List of findings compatible with VulnerabilityFinding model
        """
        findings = []

        for rule in scap_result.rules:
            rule_id = rule['rule_id']

            # Get CCIs for this rule
            ccis = cci_mapping.get(rule_id, [])

            finding = {
                'rule_id': rule_id,
                'status': rule['mapped_status'],
                'severity': self._map_severity(rule.get('severity', 'medium')),
                'cci_ids': ccis,
                'source': 'scap',
                'scan_date': scap_result.end_time.isoformat(),
                'hostname': scap_result.target_hostname,
            }

            findings.append(finding)

        return findings

    def _map_severity(self, severity: str) -> str:
        """Map SCAP severity to STIG CAT"""
        severity_map = {
            'high': 'cat1',
            'medium': 'cat2',
            'low': 'cat3',
            'critical': 'cat1',
            'info': 'cat3',
            'informational': 'cat3',
        }
        return severity_map.get(severity.lower(), 'cat2')


class RMFDocumentGenerator:
    """
    Generator for RMF authorization documents.

    Generates:
    - POA&M (Plan of Action and Milestones)
    - Test Plan
    - RAR (Risk Assessment Report)
    - Security Assessment Report
    - Authorization Package summary
    """

    def generate_poam_xlsx(
        self,
        findings: List[Dict[str, Any]],
        system_info: Dict[str, Any],
        include_milestones: bool = True
    ) -> RMFDocumentResult:
        """
        Generate POA&M in FedRAMP Excel format.

        Args:
            findings: List of open findings
            system_info: System information
            include_milestones: Whether to include milestones

        Returns:
            RMFDocumentResult with Excel content
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

            wb = Workbook()
            ws = wb.active
            ws.title = "POA&M"

            # Header styling
            header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Define columns
            columns = [
                ("A", "POA&M ID", 12),
                ("B", "Weakness Name", 40),
                ("C", "Weakness Description", 50),
                ("D", "Security Control Number", 15),
                ("E", "Point of Contact", 25),
                ("F", "Resources Required", 20),
                ("G", "Scheduled Completion Date", 20),
                ("H", "Milestone Changes", 30),
                ("I", "Status", 15),
                ("J", "Comments", 40),
                ("K", "Risk Level", 12),
                ("L", "Deviation", 15),
            ]

            if include_milestones:
                columns.extend([
                    ("M", "Milestone 1", 30),
                    ("N", "Milestone 1 Date", 15),
                    ("O", "Milestone 2", 30),
                    ("P", "Milestone 2 Date", 15),
                ])

            # Write headers
            for col_letter, header, width in columns:
                cell = ws[f"{col_letter}1"]
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = border
                ws.column_dimensions[col_letter].width = width

            # Write data
            for idx, finding in enumerate(findings, start=2):
                poam_id = finding.get('poam_id', f"POA&M-{idx:04d}")
                milestones = finding.get('milestones', [])

                row_data = [
                    poam_id,
                    finding.get('weakness_name', finding.get('title', '')),
                    finding.get('description', ''),
                    finding.get('control_id', ''),
                    finding.get('point_of_contact', ''),
                    finding.get('resources_required', ''),
                    finding.get('scheduled_completion_date', ''),
                    finding.get('milestone_changes', ''),
                    finding.get('status', 'Open'),
                    finding.get('comments', ''),
                    finding.get('risk_level', 'Moderate'),
                    finding.get('deviation', 'No'),
                ]

                if include_milestones:
                    if len(milestones) > 0:
                        row_data.extend([
                            milestones[0].get('description', ''),
                            milestones[0].get('target_date', ''),
                        ])
                    else:
                        row_data.extend(['', ''])

                    if len(milestones) > 1:
                        row_data.extend([
                            milestones[1].get('description', ''),
                            milestones[1].get('target_date', ''),
                        ])
                    else:
                        row_data.extend(['', ''])

                for col_idx, value in enumerate(row_data):
                    col_letter = chr(65 + col_idx)  # A, B, C, ...
                    cell = ws[f"{col_letter}{idx}"]
                    cell.value = value
                    cell.border = border
                    cell.alignment = Alignment(wrap_text=True, vertical='top')

            # Add system info sheet
            ws_info = wb.create_sheet("System Information")
            info_rows = [
                ("System Name", system_info.get('name', '')),
                ("System ID", system_info.get('id', '')),
                ("Authorization Type", system_info.get('authorization_type', 'FedRAMP')),
                ("Impact Level", system_info.get('impact_level', 'Moderate')),
                ("ISSO", system_info.get('isso', '')),
                ("System Owner", system_info.get('system_owner', '')),
                ("Generated Date", datetime.now().strftime('%Y-%m-%d')),
            ]

            for idx, (label, value) in enumerate(info_rows, start=1):
                ws_info[f"A{idx}"] = label
                ws_info[f"A{idx}"].font = Font(bold=True)
                ws_info[f"B{idx}"] = value

            ws_info.column_dimensions['A'].width = 25
            ws_info.column_dimensions['B'].width = 50

            # Save to bytes
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)

            return RMFDocumentResult(
                success=True,
                document_type='POA&M',
                content=buffer.getvalue(),
                filename=f"poam_{system_info.get('name', 'system')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
            )

        except ImportError:
            return RMFDocumentResult(
                success=False,
                document_type='POA&M',
                content=b'',
                filename='',
                errors=['openpyxl library not installed. Run: pip install openpyxl']
            )
        except Exception as e:
            return RMFDocumentResult(
                success=False,
                document_type='POA&M',
                content=b'',
                filename='',
                errors=[str(e)]
            )

    def generate_test_plan(
        self,
        controls: List[Dict[str, Any]],
        system_info: Dict[str, Any]
    ) -> RMFDocumentResult:
        """
        Generate Security Assessment Test Plan.

        Args:
            controls: List of controls to test
            system_info: System information

        Returns:
            RMFDocumentResult with document content
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading('Security Assessment Test Plan', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # System information
            doc.add_heading('1. System Identification', level=1)
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'

            rows = [
                ('System Name:', system_info.get('name', '')),
                ('System ID:', system_info.get('id', '')),
                ('Authorization Boundary:', system_info.get('boundary', '')),
                ('Impact Level:', system_info.get('impact_level', 'Moderate')),
                ('Assessment Date:', datetime.now().strftime('%Y-%m-%d')),
            ]

            for idx, (label, value) in enumerate(rows):
                table.rows[idx].cells[0].text = label
                table.rows[idx].cells[1].text = value

            # Assessment Scope
            doc.add_heading('2. Assessment Scope', level=1)
            doc.add_paragraph(
                f"This test plan covers the security assessment of {len(controls)} "
                f"controls for the {system_info.get('name', 'system')} system."
            )

            # Control Test Cases
            doc.add_heading('3. Control Test Cases', level=1)

            for idx, control in enumerate(controls, start=1):
                control_id = control.get('id', f'CTRL-{idx:03d}')
                control_title = control.get('title', 'Control')

                doc.add_heading(f'3.{idx} {control_id}: {control_title}', level=2)

                # Test procedure table
                test_table = doc.add_table(rows=5, cols=2)
                test_table.style = 'Table Grid'

                test_rows = [
                    ('Control ID:', control_id),
                    ('Test Objective:', control.get('objective', 'Verify control implementation')),
                    ('Test Method:', control.get('test_method', 'Examine, Interview, Test')),
                    ('Expected Results:', control.get('expected_results', 'Control is fully implemented')),
                    ('Test Procedure:', control.get('test_procedure', 'Review documentation and conduct testing')),
                ]

                for ridx, (label, value) in enumerate(test_rows):
                    test_table.rows[ridx].cells[0].text = label
                    test_table.rows[ridx].cells[1].text = value

                doc.add_paragraph()

            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return RMFDocumentResult(
                success=True,
                document_type='Test Plan',
                content=buffer.getvalue(),
                filename=f"test_plan_{system_info.get('name', 'system')}_{datetime.now().strftime('%Y%m%d')}.docx"
            )

        except ImportError:
            return RMFDocumentResult(
                success=False,
                document_type='Test Plan',
                content=b'',
                filename='',
                errors=['python-docx library not installed. Run: pip install python-docx']
            )
        except Exception as e:
            return RMFDocumentResult(
                success=False,
                document_type='Test Plan',
                content=b'',
                filename='',
                errors=[str(e)]
            )

    def generate_risk_assessment_report(
        self,
        risks: List[Dict[str, Any]],
        system_info: Dict[str, Any]
    ) -> RMFDocumentResult:
        """
        Generate Risk Assessment Report (RAR).

        Args:
            risks: List of identified risks
            system_info: System information

        Returns:
            RMFDocumentResult with document content
        """
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading('Risk Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Executive Summary
            doc.add_heading('Executive Summary', level=1)

            # Count risks by level
            risk_counts = {'High': 0, 'Moderate': 0, 'Low': 0}
            for risk in risks:
                level = risk.get('risk_level', 'Moderate')
                risk_counts[level] = risk_counts.get(level, 0) + 1

            doc.add_paragraph(
                f"This Risk Assessment Report identifies and evaluates {len(risks)} risks "
                f"for the {system_info.get('name', 'system')} system. "
                f"Of these, {risk_counts.get('High', 0)} are rated High, "
                f"{risk_counts.get('Moderate', 0)} are rated Moderate, and "
                f"{risk_counts.get('Low', 0)} are rated Low."
            )

            # System Information
            doc.add_heading('System Information', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'

            rows = [
                ('System Name:', system_info.get('name', '')),
                ('Impact Level:', system_info.get('impact_level', '')),
                ('System Owner:', system_info.get('system_owner', '')),
                ('Assessment Date:', datetime.now().strftime('%Y-%m-%d')),
            ]

            for idx, (label, value) in enumerate(rows):
                table.rows[idx].cells[0].text = label
                table.rows[idx].cells[1].text = value

            # Risk Findings
            doc.add_heading('Risk Findings', level=1)

            for idx, risk in enumerate(risks, start=1):
                risk_id = risk.get('id', f'RISK-{idx:03d}')
                risk_title = risk.get('title', 'Risk')

                doc.add_heading(f'{risk_id}: {risk_title}', level=2)

                risk_table = doc.add_table(rows=6, cols=2)
                risk_table.style = 'Table Grid'

                risk_rows = [
                    ('Risk Level:', risk.get('risk_level', 'Moderate')),
                    ('Likelihood:', risk.get('likelihood', 'Moderate')),
                    ('Impact:', risk.get('impact', 'Moderate')),
                    ('Description:', risk.get('description', '')),
                    ('Affected Controls:', ', '.join(risk.get('affected_controls', []))),
                    ('Recommended Mitigation:', risk.get('mitigation', '')),
                ]

                for ridx, (label, value) in enumerate(risk_rows):
                    risk_table.rows[ridx].cells[0].text = label
                    risk_table.rows[ridx].cells[1].text = value

                doc.add_paragraph()

            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return RMFDocumentResult(
                success=True,
                document_type='RAR',
                content=buffer.getvalue(),
                filename=f"rar_{system_info.get('name', 'system')}_{datetime.now().strftime('%Y%m%d')}.docx"
            )

        except ImportError:
            return RMFDocumentResult(
                success=False,
                document_type='RAR',
                content=b'',
                filename='',
                errors=['python-docx library not installed. Run: pip install python-docx']
            )
        except Exception as e:
            return RMFDocumentResult(
                success=False,
                document_type='RAR',
                content=b'',
                filename='',
                errors=[str(e)]
            )


class SystemScoringService:
    """
    Service for calculating aggregate system compliance scores.

    Provides:
    - System-level compliance scoring
    - Risk score calculation
    - Compliance trend tracking
    - Dashboard data generation
    """

    # Risk weights by severity category
    CAT_WEIGHTS = {
        'cat1': 10,  # High weight for critical
        'cat2': 5,   # Medium weight
        'cat3': 1,   # Low weight
    }

    def calculate_system_score(
        self,
        checklist_scores: List[Dict[str, Any]],
        system_info: Dict[str, Any]
    ) -> SystemScore:
        """
        Calculate aggregate system compliance score.

        Args:
            checklist_scores: List of ChecklistScore data
            system_info: System information

        Returns:
            SystemScore with aggregate metrics
        """
        # Aggregate counts
        total_cat1_open = 0
        total_cat1_closed = 0
        total_cat2_open = 0
        total_cat2_closed = 0
        total_cat3_open = 0
        total_cat3_closed = 0

        for score in checklist_scores:
            total_cat1_open += score.get('totalCat1Open', 0)
            total_cat1_closed += score.get('totalCat1NotAFinding', 0) + score.get('totalCat1NotApplicable', 0)
            total_cat2_open += score.get('totalCat2Open', 0)
            total_cat2_closed += score.get('totalCat2NotAFinding', 0) + score.get('totalCat2NotApplicable', 0)
            total_cat3_open += score.get('totalCat3Open', 0)
            total_cat3_closed += score.get('totalCat3NotAFinding', 0) + score.get('totalCat3NotApplicable', 0)

        # Calculate totals
        total_open = total_cat1_open + total_cat2_open + total_cat3_open
        total_closed = total_cat1_closed + total_cat2_closed + total_cat3_closed
        total_findings = total_open + total_closed

        # Calculate compliance percentage
        if total_findings > 0:
            compliance_percentage = (total_closed / total_findings) * 100
        else:
            compliance_percentage = 100.0

        # Calculate weighted risk score (higher is worse)
        risk_score = (
            total_cat1_open * self.CAT_WEIGHTS['cat1'] +
            total_cat2_open * self.CAT_WEIGHTS['cat2'] +
            total_cat3_open * self.CAT_WEIGHTS['cat3']
        )

        return SystemScore(
            system_group_id=str(system_info.get('id', '')),
            system_name=system_info.get('name', ''),
            total_checklists=len(checklist_scores),
            total_findings=total_findings,
            cat1_open=total_cat1_open,
            cat1_closed=total_cat1_closed,
            cat2_open=total_cat2_open,
            cat2_closed=total_cat2_closed,
            cat3_open=total_cat3_open,
            cat3_closed=total_cat3_closed,
            compliance_percentage=round(compliance_percentage, 2),
            risk_score=risk_score,
            last_calculated=datetime.now()
        )

    def generate_dashboard_data(
        self,
        system_scores: List[SystemScore]
    ) -> Dict[str, Any]:
        """
        Generate dashboard data from system scores.

        Args:
            system_scores: List of SystemScore objects

        Returns:
            Dashboard data structure
        """
        # Aggregate metrics
        total_systems = len(system_scores)
        total_findings = sum(s.total_findings for s in system_scores)
        total_open = sum(s.cat1_open + s.cat2_open + s.cat3_open for s in system_scores)

        # Calculate overall compliance
        total_closed = total_findings - total_open
        overall_compliance = (total_closed / total_findings * 100) if total_findings > 0 else 100.0

        # Systems by compliance level
        compliant_systems = sum(1 for s in system_scores if s.compliance_percentage >= 80)
        at_risk_systems = sum(1 for s in system_scores if 50 <= s.compliance_percentage < 80)
        non_compliant_systems = sum(1 for s in system_scores if s.compliance_percentage < 50)

        # Systems with critical findings
        critical_systems = sum(1 for s in system_scores if s.cat1_open > 0)

        # Top risks
        top_risks = sorted(system_scores, key=lambda s: s.risk_score, reverse=True)[:5]

        return {
            'summary': {
                'total_systems': total_systems,
                'total_findings': total_findings,
                'total_open': total_open,
                'overall_compliance': round(overall_compliance, 2),
                'compliant_systems': compliant_systems,
                'at_risk_systems': at_risk_systems,
                'non_compliant_systems': non_compliant_systems,
                'critical_systems': critical_systems,
            },
            'severity_breakdown': {
                'cat1_open': sum(s.cat1_open for s in system_scores),
                'cat1_closed': sum(s.cat1_closed for s in system_scores),
                'cat2_open': sum(s.cat2_open for s in system_scores),
                'cat2_closed': sum(s.cat2_closed for s in system_scores),
                'cat3_open': sum(s.cat3_open for s in system_scores),
                'cat3_closed': sum(s.cat3_closed for s in system_scores),
            },
            'top_risks': [
                {
                    'system_name': s.system_name,
                    'risk_score': s.risk_score,
                    'compliance': s.compliance_percentage,
                    'cat1_open': s.cat1_open,
                }
                for s in top_risks
            ],
            'systems': [
                {
                    'id': s.system_group_id,
                    'name': s.system_name,
                    'compliance': s.compliance_percentage,
                    'risk_score': s.risk_score,
                    'total_findings': s.total_findings,
                    'open_findings': s.cat1_open + s.cat2_open + s.cat3_open,
                    'cat1_open': s.cat1_open,
                    'cat2_open': s.cat2_open,
                    'cat3_open': s.cat3_open,
                }
                for s in system_scores
            ],
            'generated_at': datetime.now().isoformat(),
        }

    def calculate_risk_rating(self, system_score: SystemScore) -> str:
        """
        Calculate overall risk rating for a system.

        Args:
            system_score: System score data

        Returns:
            Risk rating string
        """
        # If any CAT I open, it's High risk
        if system_score.cat1_open > 0:
            return 'HIGH'

        # Based on compliance percentage and open findings
        if system_score.compliance_percentage < 50:
            return 'HIGH'
        elif system_score.compliance_percentage < 70:
            return 'MODERATE'
        elif system_score.compliance_percentage < 90:
            return 'LOW'
        else:
            return 'VERY LOW'


class AssetMetadataService:
    """
    Service for managing enhanced asset metadata.

    Provides:
    - Asset metadata extraction from scans
    - Asset inventory management
    - Technology area classification
    """

    TECHNOLOGY_AREAS = [
        'Server',
        'Workstation',
        'Network Device',
        'Database',
        'Web Server',
        'Application Server',
        'Storage',
        'Virtual Machine',
        'Container',
        'Cloud Service',
        'Mobile Device',
        'IoT Device',
        'Other',
    ]

    ASSET_CLASSIFICATIONS = [
        'Unclassified',
        'CUI',
        'Confidential',
        'Secret',
        'Top Secret',
    ]

    def extract_from_ckl(self, ckl_data: Dict[str, Any]) -> AssetMetadata:
        """
        Extract asset metadata from CKL data.

        Args:
            ckl_data: Parsed CKL data

        Returns:
            AssetMetadata object
        """
        asset_data = ckl_data.get('ASSET', {})

        # Parse IP addresses (may be comma-separated)
        ip_str = asset_data.get('HOST_IP', '')
        ip_addresses = [ip.strip() for ip in ip_str.split(',') if ip.strip()]

        # Parse MAC addresses
        mac_str = asset_data.get('HOST_MAC', '')
        mac_addresses = [mac.strip() for mac in mac_str.split(',') if mac.strip()]

        return AssetMetadata(
            hostname=asset_data.get('HOST_NAME', ''),
            ip_addresses=ip_addresses,
            mac_addresses=mac_addresses,
            fqdn=asset_data.get('HOST_FQDN', ''),
            technology_area=asset_data.get('TECH_AREA', ''),
            asset_type=asset_data.get('ASSET_TYPE', 'computing'),
            role=asset_data.get('ROLE', ''),
            operating_system=asset_data.get('TARGET_COMMENT', ''),
        )

    def extract_from_nessus(self, host_data: Dict[str, Any]) -> AssetMetadata:
        """
        Extract asset metadata from Nessus scan data.

        Args:
            host_data: Nessus host data

        Returns:
            AssetMetadata object
        """
        return AssetMetadata(
            hostname=host_data.get('name', ''),
            ip_addresses=[host_data.get('host-ip', '')] if host_data.get('host-ip') else [],
            mac_addresses=[host_data.get('mac-address', '')] if host_data.get('mac-address') else [],
            fqdn=host_data.get('host-fqdn', ''),
            operating_system=host_data.get('operating-system', ''),
            is_internet_facing=host_data.get('is-internet-facing', False),
        )

    def classify_technology_area(self, hostname: str, os_info: str) -> str:
        """
        Automatically classify technology area based on hostname and OS.

        Args:
            hostname: Asset hostname
            os_info: Operating system information

        Returns:
            Technology area classification
        """
        hostname_lower = hostname.lower()
        os_lower = os_info.lower() if os_info else ''

        # Network devices
        if any(kw in hostname_lower for kw in ['router', 'switch', 'fw', 'firewall', 'vpn']):
            return 'Network Device'

        # Databases
        if any(kw in hostname_lower for kw in ['db', 'sql', 'oracle', 'mongo', 'postgres']):
            return 'Database'

        # Web servers
        if any(kw in hostname_lower for kw in ['web', 'www', 'apache', 'nginx', 'iis']):
            return 'Web Server'

        # Workstations
        if any(kw in hostname_lower for kw in ['desktop', 'laptop', 'workstation', 'pc-']):
            return 'Workstation'

        # Check OS
        if 'windows server' in os_lower or 'linux' in os_lower:
            return 'Server'
        elif 'windows 10' in os_lower or 'windows 11' in os_lower:
            return 'Workstation'
        elif 'ios' in os_lower or 'android' in os_lower:
            return 'Mobile Device'

        return 'Server'  # Default

    def to_dict(self, metadata: AssetMetadata) -> Dict[str, Any]:
        """Convert AssetMetadata to dictionary"""
        return {
            'hostname': metadata.hostname,
            'ip_addresses': metadata.ip_addresses,
            'mac_addresses': metadata.mac_addresses,
            'fqdn': metadata.fqdn,
            'technology_area': metadata.technology_area,
            'asset_type': metadata.asset_type,
            'role': metadata.role,
            'operating_system': metadata.operating_system,
            'os_version': metadata.os_version,
            'location': metadata.location,
            'department': metadata.department,
            'system_administrator': metadata.system_administrator,
            'is_internet_facing': metadata.is_internet_facing,
            'classification': metadata.classification,
        }
