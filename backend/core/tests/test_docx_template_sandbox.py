"""Regression: docxtpl rendering of admin-uploaded Word templates must run
inside a Jinja2 sandbox so a malicious .docx can't RCE.

Two attack classes are covered:

* Classic SSTI (`{{ ''.__class__.__mro__[1].__subclasses__()[...] }}`).
* The Enterprise/PRO native-code chain from the responsible-disclosure report:
  a live Django ORM object in the template context is walked through an
  all-public attribute path to the raw DB connection
  (`audit.<manager>.all().query.get_compiler(db).connection.connection`) and
  then `load_extension` is called on an uploaded native library. The fixes are
  (1) never place a live ORM object in the context (``audit_proxy``), and
  (2) a ``HardenedReportSandbox`` that refuses attribute access on any ORM
  object as defense in depth, plus (3) rejecting native executable uploads.
"""

import io

import pytest
from docx import Document
from docxtpl import DocxTemplate
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from jinja2.exceptions import SecurityError, UndefinedError
from jinja2.sandbox import SandboxedEnvironment

from core.report_context import (
    HardenedReportSandbox,
    ReadOnlyModelProxy,
    audit_proxy,
)
from core.validators import reject_executable_content, validate_file_name


_SSTI_PAYLOAD = "{{ ''.__class__.__mro__[1].__subclasses__()[0] }}"

# The exact traversal from the disclosure report.
_ORM_GADGET = (
    "{{ audit.requirement_assessments.all().query"
    ".get_compiler(audit.requirement_assessments.db)"
    ".connection.connection }}"
)


def _docx_with(text: str) -> io.BytesIO:
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_sandbox_blocks_ssti():
    template = DocxTemplate(_docx_with(_SSTI_PAYLOAD))
    with pytest.raises(SecurityError):
        template.render({}, jinja_env=SandboxedEnvironment())


def test_unsandboxed_render_lets_ssti_through():
    """Positive control: without the sandbox the payload renders fine,
    documenting why the sandbox is required."""
    template = DocxTemplate(_docx_with(_SSTI_PAYLOAD))
    template.render({})  # would expose <class 'type'> in the rendered doc


# --- Layer 1: read-only proxy keeps live ORM objects out of the context ---


@pytest.mark.django_db
def test_proxy_exposes_only_allowlisted_fields():
    from core.models import ComplianceAssessment, Framework

    audit = ComplianceAssessment(name="My audit", version="1.0", observation="obs")
    audit.framework = Framework(name="ISO", ref_id="27001")
    proxy = audit_proxy(audit)

    assert proxy.name == "My audit"
    assert proxy.version == "1.0"
    assert proxy.framework.name == "ISO"
    assert proxy.framework.ref_id == "27001"


@pytest.mark.django_db
def test_proxy_blocks_orm_and_gadget_attributes():
    from core.models import ComplianceAssessment

    proxy = audit_proxy(ComplianceAssessment(name="x"))
    for attr in (
        "requirement_assessments",
        "evidences",
        "objects",
        "query",
        "connection",
        "path",
        "get_requirement_assessments",
        "save",
        "delete",
        "_state",
        "_meta",
    ):
        with pytest.raises(AttributeError):
            getattr(proxy, attr)


def test_proxy_relation_never_returns_live_object():
    spec = {"fields": {"name"}, "relations": {}}

    class Fake:
        name = "n"
        secret = object()

    proxy = ReadOnlyModelProxy(Fake(), spec)
    assert proxy.name == "n"
    with pytest.raises(AttributeError):
        proxy.secret  # not allowlisted


# --- Layer 2: hardened sandbox blocks ORM traversal even if it slips in ---


@pytest.mark.django_db
def test_hardened_sandbox_blocks_raw_orm_gadget():
    """A live ComplianceAssessment in the context still cannot be walked to the
    DB connection: attribute access on any ORM object is refused."""
    from core.models import ComplianceAssessment

    template = DocxTemplate(_docx_with(_ORM_GADGET))
    with pytest.raises((SecurityError, UndefinedError)):
        template.render(
            {"audit": ComplianceAssessment(name="x")},
            jinja_env=HardenedReportSandbox(),
        )


@pytest.mark.django_db
def test_proxy_plus_hardened_sandbox_fail_closed_on_gadget():
    from core.models import ComplianceAssessment

    template = DocxTemplate(_docx_with(_ORM_GADGET))
    with pytest.raises((SecurityError, UndefinedError)):
        template.render(
            {"audit": audit_proxy(ComplianceAssessment(name="x"))},
            jinja_env=HardenedReportSandbox(),
        )


@pytest.mark.django_db
def test_hardened_sandbox_still_renders_legit_fields():
    from core.models import ComplianceAssessment, Framework

    audit = ComplianceAssessment(name="Q3 audit", version="2.0")
    audit.framework = Framework(name="NIS2")
    template = DocxTemplate(
        _docx_with("{{ audit.name }}|{{ audit.framework.name }}|{{ audit.version }}")
    )
    template.render({"audit": audit_proxy(audit)}, jinja_env=HardenedReportSandbox())
    rendered = "\n".join(p.text for p in template.docx.paragraphs)
    assert "Q3 audit" in rendered
    assert "NIS2" in rendered
    assert "2.0" in rendered


# --- Layer 3: native executable uploads are rejected whatever the extension ---


@pytest.mark.parametrize(
    "magic",
    [
        b"\x7fELF\x02\x01\x01\x00",  # ELF
        b"MZ\x90\x00\x03\x00\x00\x00",  # PE
        b"\xcf\xfa\xed\xfe\x0c\x00\x00\x01",  # Mach-O 64-bit LE
        b"\xca\xfe\xba\xbe\x00\x00\x00\x02",  # Mach-O universal
    ],
)
def test_executable_content_rejected(magic):
    upload = SimpleUploadedFile("payload.zip", magic + b"rest-of-file")
    with pytest.raises(ValidationError):
        reject_executable_content(upload)
    upload.seek(0)
    with pytest.raises(ValidationError):
        validate_file_name(upload)


def test_legit_zip_upload_passes_executable_check():
    upload = SimpleUploadedFile("evidence.zip", b"PK\x03\x04rest-of-a-real-zip")
    assert reject_executable_content(upload) is upload
    # position must be restored so downstream reads/saves see the whole file
    assert upload.tell() == 0
